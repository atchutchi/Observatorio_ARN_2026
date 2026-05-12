import io
import base64

from apps.indicators.models import IndicatorCategory
from apps.dashboards.services import DashboardService
from apps.operators.models import Operator


class DocxReportGenerator:
    """Gera relatórios Word (.docx) com python-docx."""

    def __init__(self, year, quarter=None, report_type='quarterly', operator_code=None, operator_scope='all'):
        self.year = year
        self.quarter = quarter
        self.report_type = report_type
        self.operator_code = operator_code
        self.operator_scope = operator_scope

    def generate(self):
        try:
            from docx import Document
        except ImportError:
            return self._fallback()

        doc = Document()
        self._setup_styles(doc)
        context = self._build_context()

        self._add_cover_page(doc, context)
        self._add_introduction(doc, context)
        self._add_executive_summary(doc, context)
        self._add_market_shares(doc, context)
        self._add_category_sections(doc, context)
        if not context['is_operator_report']:
            self._add_hhi_section(doc, context)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _build_context(self):
        summary = DashboardService.get_summary(self.year, self.quarter, self.operator_code)
        categories = IndicatorCategory.objects.all().order_by('order')
        operators = DashboardService.get_operators(self.operator_code)

        market_shares = {}
        for market in ['mobile', 'fixed_internet', 'revenue']:
            market_shares[market] = DashboardService.get_market_share(
                self.year, self.quarter, market, self.operator_code,
            )

        category_data = {}
        for cat in categories:
            data = DashboardService.get_indicator_data(
                cat.code, self.year, self.quarter, self.operator_code,
            )
            growth = DashboardService.get_growth_rates(cat.code, self.year, self.operator_code)
            category_data[cat.code] = {
                'name': cat.name,
                'data': data,
                'growth': growth,
                'narrative': self._build_category_narrative(cat, data, growth),
            }

        hhi_mobile = DashboardService.get_hhi(self.year, 'mobile')
        hhi_internet = DashboardService.get_hhi(self.year, 'fixed_internet')

        period_label = f"Q{self.quarter} {self.year}" if self.quarter else str(self.year)
        operator_label = self._get_operator_label()

        try:
            from .chart_generator import generate_all_charts
            charts = generate_all_charts(self.year, self.quarter, self.operator_code)
        except Exception:
            charts = {}

        return {
            'year': self.year,
            'quarter': self.quarter,
            'period_label': period_label,
            'operator_scope': self.operator_scope,
            'operator_code': self.operator_code,
            'operator_label': operator_label,
            'is_operator_report': self.operator_scope != 'all',
            'operators': list(operators),
            'summary': summary,
            'market_shares': market_shares,
            'category_data': category_data,
            'hhi_mobile': hhi_mobile,
            'hhi_internet': hhi_internet,
            'categories': categories,
            'charts': charts,
        }

    def _get_operator_label(self):
        if self.operator_scope == 'others':
            return 'Outros operadores'
        if self.operator_code:
            operator = Operator.objects.filter(code=self.operator_code).first()
            if operator:
                return operator.name
        return 'Mercado geral'

    def _setup_styles(self, doc):
        from docx.shared import Pt, RGBColor

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        for level in range(1, 4):
            heading_style = doc.styles[f'Heading {level}']
            heading_style.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

    def _add_cover_page(self, doc, context):
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        for _ in range(6):
            doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('OBSERVATÓRIO DO MERCADO\nDE TELECOMUNICAÇÕES')
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
        run.bold = True

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('Guiné-Bissau')
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

        doc.add_paragraph()

        period = doc.add_paragraph()
        period.alignment = WD_ALIGN_PARAGRAPH.CENTER
        type_label = 'Relatório Trimestral' if self.report_type == 'quarterly' else 'Relatório Anual'
        run = period.add_run(f"{type_label}\n{context['period_label']}")
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        for _ in range(4):
            doc.add_paragraph()

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run('Autoridade Reguladora Nacional\nRepública da Guiné-Bissau')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

        doc.add_page_break()

    def _add_introduction(self, doc, context):
        doc.add_heading('Introdução', level=1)
        scope = (
            'o mercado das telecomunicações'
            if not context['is_operator_report']
            else f"a operadora {context['operator_label']}"
        )
        report_kind = 'trimestral' if self.report_type == 'quarterly' else 'anual'
        doc.add_paragraph(
            f"O presente relatório {report_kind} apresenta a evolução de {scope} "
            f"na Guiné-Bissau durante {context['period_label']}. "
            "A informação é organizada por indicadores regulatórios, com leitura "
            "descritiva, tabelas de suporte e gráficos de evolução."
        )
        doc.add_paragraph(
            "Os dados são tratados pela Autoridade Reguladora Nacional com base nas "
            "declarações submetidas pelos operadores licenciados e nos indicadores "
            "monitorizados pelo Observatório do Mercado."
        )

        if context['operators']:
            doc.add_heading('Operadores incluídos', level=2)
            for operator in context['operators']:
                doc.add_paragraph(operator.name, style='List Bullet')

    def _add_executive_summary(self, doc, context):
        from docx.shared import Pt

        doc.add_heading('Resumo Executivo', level=1)

        summary = context['summary']
        scope = 'o mercado' if not context['is_operator_report'] else context['operator_label']
        doc.add_paragraph(
            f"No período de {context['period_label']}, {scope} registou "
            f"{int(summary['total_subscribers']):,} assinantes "
            f"móveis no total, representando uma taxa de penetração de "
            f"{summary['penetration_rate']}%."
        )

        if summary['subscriber_change']:
            direction = 'crescimento' if summary['subscriber_change'] > 0 else 'decréscimo'
            doc.add_paragraph(
                f"Verificou-se um {direction} de {abs(summary['subscriber_change'])}% "
                f"face ao ano anterior no parque de assinantes."
            )

        doc.add_heading('Indicadores Principais', level=2)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Indicador'
        hdr[1].text = 'Valor'

        kpis = [
            ('Total Assinantes Móvel', f"{int(summary['total_subscribers']):,}"),
            ('Volume de Negócios (FCFA)', f"{int(summary['total_revenue']):,}"),
            ('Tráfego de Dados', f"{int(summary['total_data_traffic']):,}"),
            ('Taxa de Penetração', f"{summary['penetration_rate']}%"),
            ('Variação Assinantes', f"{summary['subscriber_change']}%"),
            ('Operadores Activos', str(summary['active_operators'])),
        ]

        for label, value in kpis:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value

        doc.add_paragraph()

    def _add_market_shares(self, doc, context):
        doc.add_heading('Quota de Mercado', level=1)

        market_labels = {
            'mobile': 'Mercado Móvel (Assinantes)',
            'fixed_internet': 'Internet Fixo',
            'revenue': 'Receitas',
        }

        for market_key, shares in context['market_shares'].items():
            if not shares:
                continue

            doc.add_heading(market_labels.get(market_key, market_key), level=2)
            chart_key = f'market_share_{market_key}'
            self._add_chart(doc, context['charts'].get(chart_key))

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Operador'
            hdr[1].text = 'Valor'
            hdr[2].text = 'Quota (%)'

            for s in shares:
                row = table.add_row().cells
                row[0].text = s['operator_name']
                row[1].text = f"{int(s['value']):,}"
                row[2].text = f"{s['share_pct']}%"

            doc.add_paragraph()

    def _add_category_sections(self, doc, context):
        doc.add_page_break()
        doc.add_heading('Dados por Categoria', level=1)

        for cat in context['categories']:
            cat_info = context['category_data'].get(cat.code, {})
            data = cat_info.get('data', [])
            growth = cat_info.get('growth', [])

            if not data and not growth:
                continue

            doc.add_heading(cat.name, level=2)
            narrative = cat_info.get('narrative')
            if narrative:
                doc.add_paragraph(narrative)

            chart_key = f'trends_{cat.code}'
            self._add_chart(doc, context['charts'].get(chart_key))

            if data:
                indicators = {}
                operators_set = []
                for d in data:
                    code = d['indicator_code']
                    op = d['operator_name']
                    if code not in indicators:
                        indicators[code] = {'name': d['indicator_name'], 'level': d['indicator_level'], 'values': {}}
                    if op not in operators_set:
                        operators_set.append(op)
                    if d['value'] is not None:
                        indicators[code]['values'][op] = d['value']

                if operators_set and indicators:
                    cols = 2 + len(operators_set)
                    table = doc.add_table(rows=1, cols=cols)
                    table.style = 'Light Grid Accent 1'
                    hdr = table.rows[0].cells
                    hdr[0].text = 'Cód.'
                    hdr[1].text = 'Indicador'
                    for i, op in enumerate(operators_set):
                        hdr[2 + i].text = op

                    for code, info in indicators.items():
                        row = table.add_row().cells
                        row[0].text = code
                        indent = '  ' * info['level']
                        row[1].text = f"{indent}{info['name']}"
                        for i, op in enumerate(operators_set):
                            val = info['values'].get(op)
                            row[2 + i].text = f"{int(val):,}" if val else '—'

            if growth:
                doc.add_heading('Crescimento', level=3)
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'
                hdr = table.rows[0].cells
                hdr[0].text = 'Operador'
                hdr[1].text = f'{self.year - 1}'
                hdr[2].text = f'{self.year}'
                hdr[3].text = 'Variação (%)'

                for g in growth:
                    row = table.add_row().cells
                    row[0].text = g['operator_name']
                    row[1].text = f"{int(g['previous_value']):,}"
                    row[2].text = f"{int(g['current_value']):,}"
                    row[3].text = f"{g['pct_change']}%"

            doc.add_paragraph()

    def _add_hhi_section(self, doc, context):
        doc.add_page_break()
        doc.add_heading('Índice de Concentração (HHI)', level=1)

        for label, hhi_data in [
            ('Mercado Móvel', context['hhi_mobile']),
            ('Internet Fixo', context['hhi_internet']),
        ]:
            if not hhi_data:
                continue

            doc.add_heading(label, level=2)
            if label == 'Mercado Móvel':
                self._add_chart(doc, context['charts'].get('hhi_mobile'))
            doc.add_paragraph(
                f"O índice Herfindahl-Hirschman (HHI) para o {label.lower()} em "
                f"{context['period_label']} é de {int(hhi_data['hhi']):,}, "
                f"classificado como: {hhi_data['classification']}."
            )

            doc.add_paragraph(
                'Referência: HHI < 1.500 = Competitivo | '
                '1.500 - 2.500 = Moderadamente concentrado | '
                '> 2.500 = Altamente concentrado'
            )

    def _add_chart(self, doc, chart_base64):
        if not chart_base64:
            return
        try:
            from docx.shared import Inches

            image_bytes = base64.b64decode(chart_base64)
            doc.add_picture(io.BytesIO(image_bytes), width=Inches(5.9))
        except Exception:
            return

    def _build_category_narrative(self, category, data, growth):
        if not data and not growth:
            return ''

        period = f"no {self.quarter}.º trimestre de {self.year}" if self.quarter else f"em {self.year}"
        scope = 'o mercado' if self.operator_scope == 'all' else self._get_operator_label()
        root_values = [
            item for item in data
            if item.get('indicator_level') == 0 and item.get('value') is not None
        ]
        if root_values:
            total = sum(item['value'] for item in root_values)
            return (
                f"Em {category.name}, {scope} apresentou {len(root_values)} indicadores "
                f"principais com dados {period}, totalizando {total:,.0f} nas métricas reportadas."
            )
        if growth:
            best = max(growth, key=lambda item: item.get('pct_change', 0))
            return (
                f"Em {category.name}, a maior variação anual observada {period} foi "
                f"{best['pct_change']}% em {best['operator_name']}."
            )
        return f"Em {category.name}, existem dados reportados {period} para acompanhamento."

    def _fallback(self):
        content = (
            f"Observatório Telecom GB — {self.year}\n"
            f"python-docx não disponível. Instale com: pip install python-docx\n"
        )
        return content.encode('utf-8')
